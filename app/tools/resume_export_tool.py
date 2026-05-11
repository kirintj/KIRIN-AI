import json
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.utils.chat import call_llm
from app.tools.base import BaseTool

RESUME_EXPORT_DIR = Path(__file__).resolve().parent.parent.parent / "data" / "resume_exports"

RESUME_TEMPLATES = {
    "classic": "经典简约",
    "modern": "现代专业",
    "compact": "紧凑高效",
}

RESUME_GENERATE_PROMPT = """你是一个简历生成助手。请根据以下用户信息，生成一份完整的结构化简历数据。

用户信息：
{user_info}

请以 JSON 格式输出，包含以下字段：
{{
  "name": "姓名",
  "phone": "电话",
  "email": "邮箱",
  "title": "求职意向/职位",
  "summary": "个人简介（2-3句话）",
  "skills": ["技能1", "技能2"],
  "education": [
    {{"school": "学校", "major": "专业", "degree": "学位", "period": "时间段"}}
  ],
  "experience": [
    {{"company": "公司", "position": "职位", "period": "时间段", "description": "工作描述"}}
  ],
  "projects": [
    {{"name": "项目名", "role": "角色", "period": "时间段", "description": "项目描述"}}
  ]
}}

只输出 JSON，不要其他内容。"""


def _ensure_dir(user_id: str) -> Path:
    user_dir = RESUME_EXPORT_DIR / user_id
    user_dir.mkdir(parents=True, exist_ok=True)
    return user_dir


async def generate_resume_data(user_id: str, user_info: str) -> dict:
    prompt = RESUME_GENERATE_PROMPT.format(user_info=user_info)
    raw = await call_llm(prompt, max_tokens=2000, temperature=0.3)
    return BaseTool.parse_json(raw) or {"name": "", "phone": "", "email": "", "title": "", "summary": user_info}


def _build_classic_doc(data: dict) -> Document:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    name = data.get("name", "未命名")
    title = data.get("title", "")

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(name)
    run.font.size = Pt(22)
    run.font.bold = True

    if title:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(title)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)

    contact_parts = []
    if data.get("phone"):
        contact_parts.append(f"📱 {data['phone']}")
    if data.get("email"):
        contact_parts.append(f"✉ {data['email']}")
    if contact_parts:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact.add_run(" | ".join(contact_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)

    if data.get("summary"):
        doc.add_paragraph()
        _add_section_title(doc, "个人简介")
        doc.add_paragraph(data["summary"])

    if data.get("skills"):
        doc.add_paragraph()
        _add_section_title(doc, "专业技能")
        skills_text = " | ".join(data["skills"])
        doc.add_paragraph(skills_text)

    if data.get("experience"):
        doc.add_paragraph()
        _add_section_title(doc, "工作经历")
        for exp in data["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('company', '')} - {exp.get('position', '')}")
            run.font.bold = True
            run.font.size = Pt(11)
            if exp.get("period"):
                run2 = p.add_run(f"  ({exp['period']})")
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(150, 150, 150)
            if exp.get("description"):
                doc.add_paragraph(exp["description"])

    if data.get("education"):
        doc.add_paragraph()
        _add_section_title(doc, "教育背景")
        for edu in data["education"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('school', '')} - {edu.get('major', '')} - {edu.get('degree', '')}")
            run.font.bold = True
            if edu.get("period"):
                run2 = p.add_run(f"  ({edu['period']})")
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(150, 150, 150)

    if data.get("projects"):
        doc.add_paragraph()
        _add_section_title(doc, "项目经历")
        for proj in data["projects"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{proj.get('name', '')} - {proj.get('role', '')}")
            run.font.bold = True
            if proj.get("period"):
                run2 = p.add_run(f"  ({proj['period']})")
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(150, 150, 150)
            if proj.get("description"):
                doc.add_paragraph(proj["description"])

    return doc


def _add_section_title(doc: Document, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(10, 89, 247)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(8)


def export_docx(user_id: str, data: dict, template: str = "classic") -> str:
    user_dir = _ensure_dir(user_id)
    name = data.get("name", "resume")
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"{name}_{timestamp}.docx"
    filepath = user_dir / filename

    doc = _build_classic_doc(data)
    doc.save(str(filepath))
    return str(filepath)


def export_text(data: dict) -> str:
    lines = []
    lines.append(f"# {data.get('name', '未命名')}")
    if data.get("title"):
        lines.append(f"**{data['title']}**")
    lines.append("")

    contact_parts = []
    if data.get("phone"):
        contact_parts.append(f"📱 {data['phone']}")
    if data.get("email"):
        contact_parts.append(f"✉ {data['email']}")
    if contact_parts:
        lines.append(" | ".join(contact_parts))
        lines.append("")

    if data.get("summary"):
        lines.append("## 个人简介")
        lines.append(data["summary"])
        lines.append("")

    if data.get("skills"):
        lines.append("## 专业技能")
        lines.append(" | ".join(data["skills"]))
        lines.append("")

    if data.get("experience"):
        lines.append("## 工作经历")
        for exp in data["experience"]:
            lines.append(f"### {exp.get('company', '')} - {exp.get('position', '')}")
            if exp.get("period"):
                lines.append(f"*{exp['period']}*")
            if exp.get("description"):
                lines.append(exp["description"])
            lines.append("")

    if data.get("education"):
        lines.append("## 教育背景")
        for edu in data["education"]:
            lines.append(f"- {edu.get('school', '')} | {edu.get('major', '')} | {edu.get('degree', '')} | {edu.get('period', '')}")
        lines.append("")

    if data.get("projects"):
        lines.append("## 项目经历")
        for proj in data["projects"]:
            lines.append(f"### {proj.get('name', '')} - {proj.get('role', '')}")
            if proj.get("period"):
                lines.append(f"*{proj['period']}*")
            if proj.get("description"):
                lines.append(proj["description"])
            lines.append("")

    return "\n".join(lines)


def list_exports(user_id: str) -> list[dict]:
    user_dir = _ensure_dir(user_id)
    exports = []
    for f in sorted(user_dir.glob("*.docx"), key=lambda x: x.stat().st_mtime, reverse=True):
        exports.append({
            "filename": f.name,
            "size": f.stat().st_size,
            "created_at": datetime.fromtimestamp(f.stat().st_mtime).isoformat(),
        })
    return exports
